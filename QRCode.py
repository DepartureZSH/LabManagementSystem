import qrcode
import numpy as np
import configparser
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.shared import Cm
from docx.oxml.shared import OxmlElement
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_LINE_SPACING

class QRTool:

    def __init__(self):
        config = configparser.ConfigParser()
        config.read("configs/config.ini")
        self.TempPath = config.get('Path', 'QR_Temp')
        self.outputs = config.get('Path', 'Outputs')

    def make_QRcode(self, data):
        """
        生成二维码并返回
        :param data: 二维码打包的数据
        :return: 二维码对象
        """
        qr = qrcode.QRCode(border=0)
        qr.add_data(data)
        qr.make(fit=True)
        img = qr.make_image()
        return img

    def MakeCard(self, data):
        try:
            document = Document()
            default_section = document.sections[0]
            default_section.page_width = Cm(21)
            default_section.page_height = Cm(29.7)
            default_section.top_margin = Cm(1.27)
            default_section.right_margin = Cm(1.27)
            default_section.bottom_margin = Cm(1.27)
            default_section.left_margin = Cm(1.27)

            for i in range(data.shape[0]):
                paragraph = document.add_paragraph()
                run = paragraph.add_run("{}. ".format(data.loc[i, "座位号"]))
                run.font.name = "Times New Roman"
                run.font.size = Pt(72)
                run.font.bold = True
                run = paragraph.add_run("{}".format(data.loc[i, "姓名"]))
                run.font.name = '楷体'
                run.font.size = Pt(72)
                run.font.bold = True
                run.font.element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
                try:
                    s = "姓名：{}\n".format(self.filter(data.loc[i, "姓名"]))
                    s += "学号：{}\n".format(self.filter(data.loc[i, "学号"]))
                    s += "导师：{}\n".format(self.filter(data.loc[i, "导师"]))
                    s += "实验室门牌号：{}\n".format(self.filter(data.loc[i, "实验室门牌号"]))
                    s += "座位号：{}\n".format(self.filter(data.loc[i, "座位号"]))
                    s += "工位位置：({},{})\n".format(self.filter(data.loc[i, "工位位置X"]), self.filter(data.loc[i, "工位位置Y"]))
                    img = self.make_QRcode(s)
                    run = paragraph.add_run("")
                    img.save(self.TempPath)
                    run.add_picture(self.TempPath, width=Inches(1.0), height=Inches(1.0))

                except Exception as e:
                    print(str(e))
                paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                paragraph.paragraph_format.left_indent = Cm(1.11)
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

                paragraph = document.add_paragraph()
                run = paragraph.add_run("{}：".format("学号"))
                run.font.size = Pt(42)
                run.font.name = '楷体'
                run.font.element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
                run = paragraph.add_run("{}".format(data.loc[i, "学号"]))
                run.font.name = "Times New Roman"
                run.font.size = Pt(42)
                paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                pPr = paragraph._element.get_or_add_pPr()
                ind = OxmlElement('w:ind')
                ind.set(qn('w:leftChars'), str(3 * 100))
                ind.set(qn('w:firstLineChars'), str(2 * 100))
                # if hanging_indent:
                #     ind.set(qn('w:hangingChars'), str(hanging_indent * 100))
                pPr.append(ind)

                paragraph = document.add_paragraph()
                run = paragraph.add_run("{}：{}".format("导师", data.loc[i, "导师"]))
                run.font.size = Pt(42)
                run.font.name = '楷体'
                run.font.element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
                paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                pPr = paragraph._element.get_or_add_pPr()
                ind = OxmlElement('w:ind')
                ind.set(qn('w:leftChars'), str(3 * 100))
                ind.set(qn('w:firstLineChars'), str(2 * 100))
                pPr.append(ind)

                paragraph = document.add_paragraph()
                paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                run = paragraph.add_run("{}                                               {}".format("-", "-"))
                run.font.name = "Times New Roman"
                run.font.size = Pt(42)
                pPr = paragraph._element.get_or_add_pPr()
                ind = OxmlElement('w:ind')
                ind.set(qn('w:hangingChars'), str(0.054 * 100))
                pPr.append(ind)
                paragraph.paragraph_format.left_indent = Cm(-0.02)

                paragraph = document.add_paragraph()
                run = paragraph.add_run("{}".format("请沿着上边标记折叠"))
                run.font.size = Pt(11)
                run.font.name = '楷体'
                run.font.element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
                paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)

                # document.add_page_break()
                # 添加分页
                document.add_page_break()
            document.save(self.outputs + "学生工位牌.docx")
            return True
        except Exception as e:
            print(str(e))
            return False

    def filter(self, data):
        Type = type(data)
        if str(data) == "nan":
            return ""
        elif Type == str:
            return data
        elif Type == int or Type == np.int64 or Type == np.int32:
            return str(data)
        else:
            return str(int(data))
